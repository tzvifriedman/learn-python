import csv
import os
from docx import Document

f = open("StudentFiles/students2.csv")
csv_f = csv.reader(f)


for row in csv_f:

    template = Document("StudentFiles/template.docx")
    for para in template.paragraphs:
        if 'Last' in para.text:
            try:
                para.text = row[0].strip()
                try:
                    os.mkdir("./StudentFiles/Graduating Class " + row[2])
                except: 
                    pass
                try: 
                    os.mkdir("./StudentFiles/Graduating Class " + row[2] + "/" + para.text)
                except: 
                    pass
                try: 
                    os.mkdir("./StudentFiles/Graduating Class " + row[2] + "/" + para.text + "/" + row[1])
                except:
                    pass
                
                template.save("./StudentFiles/Graduating Class " + row[2] + "/" + para.text + "/" + para.text + ".docx")
            except:
                print("Not a file")

# for para in template.paragraphs:
#     if 'Last' in para.text:
#         para.text = "New names"
#         template.save(f"/Users/tzvi.friedman/Downloads/edited.docx")
